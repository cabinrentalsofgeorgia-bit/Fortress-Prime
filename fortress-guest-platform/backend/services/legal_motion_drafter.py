"""
Legal motion drafter — produces court-ready DOCX artifacts for deadline-driven
extension motions in the sovereign legal workflow.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 12,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    indent: float = 0.0,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"


def generate_motion_extension_docx(
    *,
    case_number: str,
    case_name: str,
    court: str,
    judge: str,
    jurisdiction: str,
    deadline_date: str,
    deadline_type: str,
    days_remaining: int,
    supporting_context: str | None = None,
) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    _add_paragraph(doc, f"IN THE {court}".upper(), bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, jurisdiction.upper(), bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"Civil Action No. {case_number}", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(doc, "")
    _add_paragraph(doc, case_name, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")
    _add_paragraph(doc, "MOTION FOR EXTENSION OF TIME", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "")

    today = datetime.utcnow().strftime("%B %d, %Y")
    _add_paragraph(
        doc,
        f"COMES NOW the Defendant and respectfully moves this Court, with the matter presently before "
        f"{judge}, for a limited extension of time regarding the {deadline_type.lower()} presently set "
        f"for {deadline_date}. This request is submitted on {today}.",
    )
    _add_paragraph(doc, "")
    _add_paragraph(doc, "In support of this Motion, Defendant states as follows:", bold=True)
    _add_paragraph(doc, "")
    _add_paragraph(
        doc,
        f"1. The operative deadline for {deadline_type} presently falls on {deadline_date}, leaving "
        f"{days_remaining} day(s) remaining to complete responsive work in a matter involving active legal strategy review.",
    )
    _add_paragraph(
        doc,
        "2. Additional time is necessary to finalize responsive materials, verify the complete factual record, "
        "and ensure submissions comply with the Court's scheduling expectations and local practice.",
    )
    _add_paragraph(
        doc,
        f"3. This request is made in good faith and not for purposes of delay, but to promote an orderly presentation "
        f"of the issues before the Court under Judge {judge}'s protocols within the {jurisdiction}.",
    )
    if supporting_context:
        _add_paragraph(
            doc,
            f"4. Additional context supporting this request: {supporting_context.strip()}",
        )

    _add_paragraph(doc, "")
    _add_paragraph(doc, "WHEREFORE, Defendant respectfully requests that the Court:", bold=True)
    _add_paragraph(doc, "a. grant a reasonable extension of time for the identified deadline;", indent=0.25)
    _add_paragraph(doc, "b. accept any resulting filing as timely if submitted within the extended period; and", indent=0.25)
    _add_paragraph(doc, "c. award such other and further relief as the Court deems just and proper.", indent=0.25)

    _add_paragraph(doc, "")
    _add_paragraph(doc, f"Respectfully submitted this {today}.")
    _add_paragraph(doc, "")
    _add_paragraph(doc, "Gary Knight, Manager", bold=True)
    _add_paragraph(doc, "Cabin Rentals of Georgia, LLC")
    _add_paragraph(doc, "Pro Se Defendant")
    _add_paragraph(doc, "cabin.rentals.of.georgia@gmail.com")

    _add_paragraph(doc, "")
    _add_paragraph(doc, "CERTIFICATE OF SERVICE", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(
        doc,
        "I certify that I have served a true and correct copy of the foregoing Motion for Extension of Time "
        "upon all counsel of record in accordance with the Court's rules and applicable service requirements.",
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def motion_extension_filename(*, case_number: str) -> str:
    safe_case = "".join(ch if ch.isalnum() or ch in {"-", "_"} else "_" for ch in case_number.strip())
    return f"Motion_Extension_of_Time_{safe_case}.docx"

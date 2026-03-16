"""
FORTRESS PRIME — LEGAL DOCUMENT GENERATOR
==========================================
Generates court-formatted DOCX pleadings (Answer and Affirmative Defenses)
for Georgia Superior Court using the Legal Council's consensus output.

Library: python-docx 1.2.0

Output format:
    - Caption block (Georgia Superior Court formatting)
    - Case number, parties, judge
    - Numbered affirmative defenses derived from Council consensus
    - Prayer for relief
    - Verification / signature block
    - Certificate of Service
"""

import io
import re
import structlog
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = structlog.get_logger()


def _add_centered(doc: Document, text: str, bold: bool = False, size: int = 12, caps: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper() if caps else text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def _add_left(doc: Document, text: str, bold: bool = False, size: int = 12, indent: float = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return p


def _add_numbered(doc: Document, number: int, text: str, size: int = 12, indent: float = 0.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run(f"{number}. ")
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run2 = p.add_run(text)
    run2.font.size = Pt(size)
    run2.font.name = "Times New Roman"
    return p


def _add_blank_line(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(12)
    run.font.name = "Times New Roman"


def _extract_case_meta(case_brief: str) -> dict:
    """Extract case metadata from the brief text using pattern matching."""
    meta = {
        "case_number": "SUV2026000013",
        "court": "SUPERIOR COURT OF FANNIN COUNTY",
        "circuit": "APPALACHIAN JUDICIAL CIRCUIT",
        "state": "STATE OF GEORGIA",
        "plaintiff": "GENERALI GLOBAL ASSISTANCE, INC.",
        "defendant": "CABIN RENTALS OF GEORGIA, LLC",
        "judge": "Honorable J. David Stuart",
    }

    cn_match = re.search(r"CASE\s*NUMBER[:\s]*(\S+)", case_brief, re.IGNORECASE)
    if cn_match:
        meta["case_number"] = cn_match.group(1)

    judge_match = re.search(r"JUDGE[:\s]*(.+?)(?:\n|$)", case_brief, re.IGNORECASE)
    if judge_match:
        meta["judge"] = judge_match.group(1).strip()

    entity_suffix = r"(?:INC|LLC|CORP|LTD|L\.?P)\.?"
    plaintiff_match = re.search(
        rf"PLAINTIFF[:\s]*([A-Z][A-Z\s,.'&]+?{entity_suffix})\b",
        case_brief, re.IGNORECASE,
    )
    if plaintiff_match:
        meta["plaintiff"] = plaintiff_match.group(1).strip().rstrip(".").upper()

    defendant_match = re.search(
        rf"DEFENDANT[:\s]*([A-Z][A-Z\s,.'&]+?{entity_suffix})\b",
        case_brief, re.IGNORECASE,
    )
    if defendant_match:
        raw = defendant_match.group(1).strip().rstrip(".")
        raw = re.sub(r'\s*\(.*\)', '', raw)
        raw = re.sub(r'\s*d/b/a.*', '', raw, flags=re.IGNORECASE)
        meta["defendant"] = raw.upper()

    return meta


def generate_answer_and_defenses(
    case_brief: str,
    consensus: dict[str, Any],
) -> bytes:
    """
    Generate a Georgia Superior Court Answer and Affirmative Defenses DOCX.

    Returns the raw bytes of the DOCX file.
    """
    meta = _extract_case_meta(case_brief)
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Caption Block ──────────────────────────────────────────────────

    _add_centered(doc, f"IN THE {meta['court']}", bold=True, size=13, caps=True)
    _add_centered(doc, meta["circuit"], bold=True, size=12, caps=True)
    _add_centered(doc, meta["state"], bold=True, size=12, caps=True)
    _add_blank_line(doc)

    _add_left(doc, f"{meta['plaintiff']},", bold=True)
    _add_left(doc, "Plaintiff,", indent=0.5)
    _add_blank_line(doc)
    _add_centered(doc, "v.", bold=True, size=12)
    _add_blank_line(doc)
    _add_left(doc, f"{meta['defendant']},", bold=True)
    _add_left(doc, "Defendant.", indent=0.5)

    p_case = doc.add_paragraph()
    p_case.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_case = p_case.add_run(f"Civil Action No. {meta['case_number']}")
    run_case.bold = True
    run_case.font.size = Pt(12)
    run_case.font.name = "Times New Roman"

    _add_blank_line(doc)

    # ── Horizontal rule (simulated) ────────────────────────────────────
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rule = p_rule.add_run("_" * 60)
    run_rule.font.size = Pt(10)
    run_rule.font.name = "Times New Roman"
    run_rule.font.color.rgb = RGBColor(128, 128, 128)

    _add_blank_line(doc)

    # ── Title ──────────────────────────────────────────────────────────

    _add_centered(
        doc,
        "DEFENDANT'S ANSWER, AFFIRMATIVE DEFENSES,",
        bold=True, size=14, caps=True,
    )
    _add_centered(
        doc,
        "AND ADDITIONAL DEFENSES",
        bold=True, size=14, caps=True,
    )
    _add_blank_line(doc)

    # ── Introduction ───────────────────────────────────────────────────

    signal = consensus.get("consensus_signal", "DEFENSE")
    conviction = consensus.get("consensus_conviction", 0)
    conviction_pct = round(conviction * 100) if isinstance(conviction, (int, float)) else 0

    _add_left(
        doc,
        f"COMES NOW the Defendant, {meta['defendant']}, pro se, and for its "
        f"Answer to the Complaint filed herein, states as follows:",
    )
    _add_blank_line(doc)

    # ── General Denials ────────────────────────────────────────────────

    _add_centered(doc, "GENERAL RESPONSE", bold=True, size=13)
    _add_blank_line(doc)

    _add_numbered(doc, 1,
        "Defendant denies each and every allegation of the Complaint not "
        "specifically admitted herein.")
    _add_numbered(doc, 2,
        "Defendant demands strict proof of all allegations set forth in the Complaint.")
    _add_numbered(doc, 3,
        "To the extent any allegation of the Complaint is not specifically addressed "
        "herein, such allegation is denied.")
    _add_blank_line(doc)

    # ── Affirmative Defenses ───────────────────────────────────────────

    _add_centered(doc, "AFFIRMATIVE DEFENSES", bold=True, size=13)
    _add_blank_line(doc)

    defenses = consensus.get("top_defense_arguments", [])
    if not defenses:
        defenses = [
            "Lack of privity of contract between Plaintiff and Defendant.",
            "Statute of limitations bars some or all of Plaintiff's claims.",
            "Failure to state a claim upon which relief can be granted.",
        ]

    for i, defense in enumerate(defenses, start=1):
        defense_clean = defense.strip()
        if not defense_clean.endswith("."):
            defense_clean += "."
        _add_left(doc, f"DEFENSE {i}", bold=True, size=12)
        _add_left(doc, defense_clean, indent=0.5)
        _add_blank_line(doc)

    # ── Risk Acknowledgment (internal, marked as such) ─────────────────

    risks = consensus.get("top_risk_factors", [])
    if risks:
        _add_centered(doc, "MATTERS REQUIRING ATTENTION", bold=True, size=13)
        p_note = doc.add_paragraph()
        p_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_note = p_note.add_run(
            "[INTERNAL NOTE — NOT FOR FILING: The following risk factors were "
            "identified by the Legal Council and should be addressed in trial preparation.]"
        )
        run_note.italic = True
        run_note.font.size = Pt(10)
        run_note.font.name = "Times New Roman"
        run_note.font.color.rgb = RGBColor(180, 0, 0)

        for i, risk in enumerate(risks, start=1):
            risk_clean = risk.strip()
            if not risk_clean.endswith("."):
                risk_clean += "."
            _add_numbered(doc, i, risk_clean, size=10)
        _add_blank_line(doc)

    # ── Prayer for Relief ──────────────────────────────────────────────

    _add_centered(doc, "PRAYER FOR RELIEF", bold=True, size=13)
    _add_blank_line(doc)
    _add_left(doc, "WHEREFORE, Defendant respectfully prays that this Honorable Court:")
    _add_blank_line(doc)
    _add_numbered(doc, 1, "Dismiss Plaintiff's Complaint in its entirety with prejudice;")
    _add_numbered(doc, 2, "Award Defendant its costs of defense incurred herein;")
    _add_numbered(doc, 3, "Grant such other and further relief as this Court deems just and proper.")
    _add_blank_line(doc)

    # ── Jury Demand ────────────────────────────────────────────────────

    _add_centered(doc, "DEMAND FOR JURY TRIAL", bold=True, size=13)
    _add_blank_line(doc)
    _add_left(
        doc,
        "Defendant hereby demands trial by jury on all issues so triable.",
    )
    _add_blank_line(doc)
    _add_blank_line(doc)

    # ── Signature Block ────────────────────────────────────────────────

    _add_left(doc, f"Respectfully submitted this {datetime.now().strftime('%d')} day of "
              f"{datetime.now().strftime('%B, %Y')}.")
    _add_blank_line(doc)
    _add_blank_line(doc)
    _add_left(doc, "____________________________________")
    _add_left(doc, "Gary Knight, Manager", bold=True)
    _add_left(doc, "Cabin Rentals of Georgia, LLC")
    _add_left(doc, "Pro Se Defendant")
    _add_left(doc, "PO Box 982")
    _add_left(doc, "Morganton, GA 30560")
    _add_left(doc, "(678) 549-3680")
    _add_left(doc, "cabin.rentals.of.georgia@gmail.com")
    _add_blank_line(doc)

    # ── Certificate of Service ─────────────────────────────────────────

    p_rule2 = doc.add_paragraph()
    p_rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_rule2 = p_rule2.add_run("_" * 60)
    run_rule2.font.size = Pt(10)
    run_rule2.font.name = "Times New Roman"
    run_rule2.font.color.rgb = RGBColor(128, 128, 128)

    _add_blank_line(doc)
    _add_centered(doc, "CERTIFICATE OF SERVICE", bold=True, size=13)
    _add_blank_line(doc)
    _add_left(
        doc,
        "I hereby certify that I have this day served a true and correct copy of the "
        "foregoing DEFENDANT'S ANSWER, AFFIRMATIVE DEFENSES, AND ADDITIONAL DEFENSES "
        "upon the Plaintiff's counsel of record by depositing the same in the United States Mail, "
        "first-class postage prepaid, addressed as follows:",
    )
    _add_blank_line(doc)
    _add_left(doc, "J. David Stuart, Esq.", indent=0.5)
    _add_left(doc, "Stuart Attorneys", indent=0.5)
    _add_left(doc, "PO Box 1567", indent=0.5)
    _add_left(doc, "Smyrna, GA 30081", indent=0.5)
    _add_blank_line(doc)
    _add_left(doc, f"This {datetime.now().strftime('%d')} day of {datetime.now().strftime('%B, %Y')}.")
    _add_blank_line(doc)
    _add_left(doc, "____________________________________")
    _add_left(doc, "Gary Knight, Pro Se Defendant")

    # ── Council Metadata Footer ────────────────────────────────────────

    _add_blank_line(doc)
    _add_blank_line(doc)
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_text = (
        f"[Generated by Fortress Prime Legal Council — "
        f"Signal: {signal} | Conviction: {conviction_pct}% | "
        f"Defenses: {consensus.get('defense_count', 0)} | "
        f"Voters: {consensus.get('total_voters', 0)} | "
        f"{datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}]"
    )
    run_meta = p_meta.add_run(meta_text)
    run_meta.italic = True
    run_meta.font.size = Pt(8)
    run_meta.font.name = "Times New Roman"
    run_meta.font.color.rgb = RGBColor(128, 128, 128)

    # ── Serialize to bytes ─────────────────────────────────────────────

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    logger.info(
        "legal_docgen_complete",
        case_number=meta["case_number"],
        signal=signal,
        num_defenses=len(defenses),
        num_risks=len(risks),
    )

    return buf.getvalue()
